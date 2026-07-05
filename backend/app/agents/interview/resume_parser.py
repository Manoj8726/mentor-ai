import os
import fitz
import docx
import logging

logger = logging.getLogger("app.agents.interview.resume_parser")


class ResumeParser:
    """
    Utility parser handling text extraction from PDF and DOCX resume formats.
    """

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Detects file extension and extracts raw text string.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found at: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return ResumeParser._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return ResumeParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: '{ext}'. Only PDF and DOCX files are supported.")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        logger.info(f"Extracting PDF text from: {file_path}")
        text_content = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_content.append(page.get_text())
            return "\n".join(text_content)
        except Exception as e:
            logger.error(f"Failed parsing PDF file {file_path}: {e}")
            raise ValueError(f"Error reading PDF content: {str(e)}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        logger.info(f"Extracting DOCX text from: {file_path}")
        try:
            doc = docx.Document(file_path)
            paragraphs_text = [p.text for p in doc.paragraphs]
            # Include table contents if available to catch contact details/skills grids
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs_text.append(cell.text)
            return "\n".join(paragraphs_text)
        except Exception as e:
            logger.error(f"Failed parsing DOCX file {file_path}: {e}")
            raise ValueError(f"Error reading DOCX content: {str(e)}")
